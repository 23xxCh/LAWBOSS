"""生成出海法盾 CrossGuard 商业计划书 DOCX"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 全局样式设置 =====
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)

def add_table(headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1A3A6E',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_para(text, bold=False, size=12, align=None, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('出海法盾 CrossGuard')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('跨境电商智能合规审查平台')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4A, 0x6A, 0x8E)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('商 业 计 划 书')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    '项目版本：v0.2.0',
    '编制日期：2026年4月',
    '密级：机密',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
doc.add_heading('目  录', level=1)
toc_items = [
    ('一、执行摘要', 3),
    ('二、项目背景与市场分析', 5),
    ('  2.1 行业背景', 5),
    ('  2.2 市场规模', 6),
    ('  2.3 竞争分析', 7),
    ('  2.4 目标客户', 8),
    ('三、产品与技术', 9),
    ('  3.1 产品定位', 9),
    ('  3.2 核心功能', 10),
    ('  3.3 技术架构', 12),
    ('  3.4 技术壁垒与护城河', 13),
    ('  3.5 产品路线图', 14),
    ('四、商业模式', 15),
    ('  4.1 盈利模式', 15),
    ('  4.2 定价策略', 16),
    ('  4.3 单位经济模型', 17),
    ('五、运营与增长策略', 18),
    ('  5.1 获客策略', 18),
    ('  5.2 增长飞轮', 19),
    ('  5.3 合作生态', 20),
    ('六、团队与组织', 21),
    ('七、财务预测', 22),
    ('八、风险分析与应对', 24),
    ('九、社会价值', 26),
    ('十、附录', 27),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{item}')
    run.font.size = Pt(12)
    if not item.startswith('  '):
        run.bold = True

doc.add_page_break()

# ============================================================
# 一、执行摘要
# ============================================================
doc.add_heading('一、执行摘要', level=1)

add_para('出海法盾（CrossGuard）是面向跨境电商卖家的智能合规审查SaaS平台，致力于解决中国出海电商在多国市场下面临的法规合规痛点。平台采用"两级检测流水线"架构——关键词快速初筛 + AI语义深度复筛，能够自动识别产品描述中的医疗宣称、绝对化用语、虚假广告、缺失标签、禁用成分、隐含违规等6类违规内容，秒级生成合规报告并提供修改建议和合规版本。', indent=True)

add_para('核心创新点：', bold=True)
add_para('（1）两级检测流水线：关键词确定性初筛（高召回）+ LLM语义复筛（低误报），兼顾速度与精度；', indent=True)
add_para('（2）数据飞轮：用户反馈（误报/漏报标记）→ 统计分析 → 规则优化建议 → 人工审核 → 规则更新 → 精度提升，形成越用越准的正向循环；', indent=True)
add_para('（3）平台对接闭环：Amazon SP-API / Shopee Open API 自动拉取Listing → 合规检测 → 高风险告警 → 合规版本回写，从"工具"升级为"平台"；', indent=True)
add_para('（4）5市场覆盖：欧盟、美国、新加坡、泰国、马来西亚，涵盖欧盟EC 1223/2009、美国FD&C Act、东盟ACD等主流法规体系。', indent=True)

add_para('市场机会：中国跨境电商出口额2025年预计突破2.5万亿元，其中东南亚市场增速超25%。因产品描述违规导致的下架、罚款、诉讼每年造成数十亿元损失，合规审查已成为出海电商的刚需。', indent=True)

add_para('商业模式：SaaS订阅制（免费版/专业版/企业版/定制版四档定价），辅以API调用计费和合规培训增值服务。预计第3年实现盈亏平衡，第5年ARR突破5000万元。', indent=True)

# ============================================================
# 二、项目背景与市场分析
# ============================================================
doc.add_heading('二、项目背景与市场分析', level=1)

doc.add_heading('2.1 行业背景', level=2)

add_para('跨境电商正处于从"野蛮生长"向"合规经营"转型的关键期。近年来，各国监管机构对电商产品描述合规的执法力度显著加强：', indent=True)

add_table(
    ['时间', '事件', '影响'],
    [
        ['2023年', '欧盟CPNP通报召回超2000款化妆品', '违规产品强制下架+品牌声誉损失'],
        ['2023年', '美国FTC对虚假广告开出$500K罚单', '单次罚款可达百万美元级'],
        ['2024年', 'FDA对膳食补充剂发出47封警告信', '产品强制下架+进口禁令'],
        ['2024年', '新加坡HSA加强电商化妆品监管', '东南亚执法趋严'],
        ['2025年', '东盟ACD统一化妆品监管框架生效', '合规门槛提高但规则趋同'],
    ],
    col_widths=[2.5, 5.5, 5]
)

add_para('核心痛点：出海卖家同时面对5个以上目标市场、数十个产品类别、数百条法规条款，人工审查一条产品描述需30-60分钟且极易遗漏，而违规成本极高（下架损失+罚款+诉讼费+品牌损伤）。', indent=True)

doc.add_heading('2.2 市场规模', level=2)

add_table(
    ['市场', '2025年规模(估算)', '年增长率', '合规审查需求特征'],
    [
        ['中国跨境电商出口总额', '2.5万亿元', '15%', '合规是出海刚需'],
        ['跨境电商SaaS市场', '350亿元', '25%', '工具化渗透率快速提升'],
        ['合规审查细分市场', '30亿元', '35%', 'AI合规审查处于早期爆发阶段'],
        ['东南亚电商市场', '1.2万亿元', '25%+', '增速最快，合规意识觉醒'],
    ],
    col_widths=[4, 3, 2, 5]
)

add_para('TAM/SAM/SOM分析：', bold=True)
add_para('• TAM（总可寻址市场）：中国跨境电商卖家合规审查相关支出 ≈ 30亿元/年', indent=True)
add_para('• SAM（可服务市场）：采用SaaS工具进行合规审查的卖家支出 ≈ 8亿元/年', indent=True)
add_para('• SOM（可获得市场）：CrossGuard前3年可触达的市场份额 ≈ 1.2亿元/年', indent=True)

doc.add_heading('2.3 竞争分析', level=2)

add_table(
    ['维度', 'CrossGuard', '传统律所/法务', 'Compliance.ai', '卖家自查'],
    [
        ['检测速度', '秒级/条', '小时-天级', '分钟级', '30-60分钟/条'],
        ['覆盖市场', '5大市场', '1-2个市场', '2-3个市场', '1个市场'],
        ['AI语义检测', '两级流水线', '人工判断', '规则引擎', '无'],
        ['数据飞轮', '反馈→优化→提升', '无', '无', '无'],
        ['平台对接', 'Amazon/Shopee API', '无', '部分', '无'],
        ['成本/条', '¥0.1-1', '¥50-200', '¥2-5', '¥10-30(人力)'],
        ['合规版本', '自动生成', '人工改写', '无', '人工改写'],
    ],
    col_widths=[2.5, 2.5, 2.5, 2.5, 2.5]
)

add_para('竞争优势总结：CrossGuard的核心壁垒不在于LLM调用（这是商品化能力），而在于三层独占资产的叠加：', indent=True)
add_para('（1）法规知识库：5市场×多类别的禁用词库、法规数据、替换映射、违规案例，需持续专业维护；', indent=True)
add_para('（2）用户反馈数据：数据飞轮积累的误报/漏报标注数据，是优化规则精度的唯一来源，竞对无法复制；', indent=True)
add_para('（3）平台API集成：与Amazon/Shopee的深度对接，形成使用粘性和迁移成本。', indent=True)

doc.add_heading('2.4 目标客户', level=2)

add_table(
    ['客户画像', '规模估算', '核心需求', '付费意愿', '获客渠道'],
    [
        ['中小卖家(年GMV<5000万)', '50万+', '快速检测+合规版本', '¥99-299/月', '社群/内容/平台'],
        ['大卖/品牌卖家(年GMV>5000万)', '5万+', '批量检测+巡检+API', '¥999-4999/月', '直销/渠道合作'],
        ['电商服务商/代运营', '1万+', '白标API+客户管理', '¥4999+/月', 'BD合作'],
        ['合规/法务机构', '5000+', '辅助审查+报告导出', '¥999/月', '行业会议/推荐'],
    ],
    col_widths=[3.5, 2, 3, 2.5, 2.5]
)

# ============================================================
# 三、产品与技术
# ============================================================
doc.add_heading('三、产品与技术', level=1)

doc.add_heading('3.1 产品定位', level=2)

add_para('CrossGuard定位为"跨境电商合规基础设施"——从单点工具出发，通过平台对接和数据飞轮，逐步演化为出海合规生态的核心节点。', indent=True)

add_table(
    ['演进阶段', '定位', '核心能力', '护城河'],
    [
        ['V1(当前)', '合规检测工具', '两级检测+合规版本+报告导出', '法规知识库'],
        ['V2(6个月)', '合规管理平台', '+平台对接+巡检+数据飞轮', '知识库+反馈数据+平台绑定'],
        ['V3(12个月)', '合规基础设施', '+多平台对接+规则市场+合规认证', '数据网络效应+生态锁定'],
    ],
    col_widths=[2.5, 3, 4, 4]
)

doc.add_heading('3.2 核心功能', level=2)

add_para('3.2.1 两级检测流水线', bold=True)
add_para('第一级——关键词初筛（5个检测器）：', indent=True)
add_table(
    ['检测器', '违规类型', '评分', '数据来源'],
    [
        ['MedicalClaimChecker', '医疗宣称', '25分/项', 'EU/US/SEA禁用词库(382词)'],
        ['AbsoluteTermChecker', '绝对化用语', '15分/项', 'absolute_terms.txt(75词)'],
        ['FalseAdChecker', '虚假广告', '15分/项', '8组市场×类别正则模式'],
        ['MissingLabelChecker', '缺失标签', '10分/项', '7组市场×类别标签关键词'],
        ['BannedIngredientChecker', '禁用成分', '30分/项', 'EU/US/SEA成分库(109种)'],
    ],
    col_widths=[3.5, 2.5, 2, 5]
)

add_para('第二级——AI语义复筛（AISemanticChecker）：', indent=True)
add_para('基于LLM（GPT-4o/DeepSeek/Ollama）的5维度深度检测：隐含医疗宣称、误导性功效宣称、不当对比与贬低、文化/法律语境违规、成分暗示违规。仅报告置信度>0.7的违规，避免误报。', indent=True)

add_para('3.2.2 合规版本生成', bold=True)
add_para('上下文感知替换：只替换被标记为违规的词汇，而非全文替换。例如"治疗"被标记为医疗宣称违规时替换为"舒缓"，但"治疗"出现在非违规语境时不替换。EU市场30组、US市场20组、SEA市场26组替换映射。', indent=True)

add_para('3.2.3 图片合规检测', bold=True)
add_para('Tesseract OCR提取图片文字 → 文本合规检测。支持PNG/JPG/WEBP格式，自动识别Amazon Listing主图中的违规文字。', indent=True)

add_para('3.2.4 电商平台对接', bold=True)
add_table(
    ['平台', 'API', '能力', '类别映射'],
    [
        ['Amazon', 'SP-API', '拉取Catalog Items/更新Listing', '10组类别映射'],
        ['Shopee', 'Open API v2', '拉取商品列表/更新描述', '5组类别映射'],
    ],
    col_widths=[2, 3, 4, 3]
)

add_para('3.2.5 合规巡检引擎', bold=True)
add_para('自动流程：平台API拉取Listing → 逐条合规检测 → 高风险项生成告警 → Webhook通知（飞书/钉钉/Slack） → 巡检结果存储。支持手动触发和定期调度。', indent=True)

add_para('3.2.6 数据飞轮', bold=True)
add_para('核心闭环：用户反馈（误报/漏报/正确标记）→ 精度指标计算（准确率/误报率/漏报率）→ 规则优化建议自动生成（误报≥3次建议移除、漏报≥2次建议补充）→ 人工审核 → 规则更新 → 精度提升。', indent=True)

add_para('3.2.7 其他功能', bold=True)
add_para('• 批量检测（最多100条/次）• PDF报告导出 • 必需标签/认证查询 • JWT+RBAC认证 • 风险评分仪表盘 • Diff对比可视化', indent=True)

doc.add_heading('3.3 技术架构', level=2)

add_table(
    ['层级', '技术选型', '选型理由'],
    [
        ['后端框架', 'FastAPI + Pydantic v2', '异步高性能、自动OpenAPI文档、类型安全'],
        ['ORM', 'SQLAlchemy 2.0 + PostgreSQL/SQLite', '生产PostgreSQL+开发SQLite双模式'],
        ['AI检测', 'LLM API (GPT-4o/DeepSeek/Ollama)', '多后端支持、可切换、本地部署可选'],
        ['OCR', 'Tesseract + Pillow', '开源成熟、支持中英文+东南亚语言'],
        ['认证', 'JWT (python-jose) + bcrypt', '无状态认证、工业级密码哈希'],
        ['报告导出', 'Jinja2 + WeasyPrint', 'HTML模板→PDF、样式可控'],
        ['前端', 'React 19 + TypeScript + Ant Design 6', '类型安全、企业级UI组件'],
        ['容器化', 'Docker + docker-compose', '一键部署、环境一致性'],
        ['CI/CD', 'GitHub Actions', '自动化测试+构建+部署'],
    ],
    col_widths=[2.5, 4, 6]
)

add_para('架构设计原则：', bold=True)
add_para('• 可插拔检测器：BaseChecker抽象基类，新增检测器只需继承并实现check()方法；', indent=True)
add_para('• 可扩展平台：BasePlatformClient抽象基类，新增平台只需实现4个方法；', indent=True)
add_para('• 词边界匹配：英文使用\\b正则边界，避免子串误报（painting不匹配pain）；', indent=True)
add_para('• 环境变量驱动：所有外部依赖（数据库/LLM/平台API）通过环境变量配置，12-Factor App规范。', indent=True)

doc.add_heading('3.4 技术壁垒与护城河', level=2)

add_table(
    ['壁垒层级', '资产', '构建方式', '竞对复制难度'],
    [
        ['L1-数据资产', '法规知识库(382+禁用词/109禁用成分/76替换映射/4违规案例)', '专业法务持续维护', '中等-需专业投入'],
        ['L2-反馈数据', '用户误报/漏报标注数据', '数据飞轮自动积累', '高-时间依赖+用户锁定'],
        ['L3-平台集成', 'Amazon/Shopee API对接+类别映射', '工程开发+平台审核', '高-需平台授权'],
        ['L4-网络效应', '更多用户→更多反馈→更准规则→更好产品', '飞轮自增强', '极高-需同时具备L1-L3'],
    ],
    col_widths=[2, 4, 3, 3]
)

add_para('关键认知：LLM能力是商品化的（任何人都能接入GPT-4），但法规知识库+用户反馈数据+平台API集成是独占资产。CrossGuard的护城河不在AI算法本身，而在AI与领域知识的深度融合。', indent=True)

doc.add_heading('3.5 产品路线图', level=2)

add_table(
    ['阶段', '时间', '核心交付', '关键指标'],
    [
        ['MVP(当前)', 'V0.2.0', '5市场两级检测+合规版本+PDF导出+认证+数据飞轮+平台对接', '核心功能可用'],
        ['Phase 1', '1-3个月', 'Lazada/TikTok Shop对接+定期巡检调度+规则管理后台', '平台对接≥3个'],
        ['Phase 2', '3-6个月', '多语言UI(英/中/泰)+团队协作+白标API+合规培训模块', '付费用户≥100'],
        ['Phase 3', '6-12个月', '规则市场(用户贡献规则)+合规认证服务+行业解决方案', 'ARR≥500万'],
    ],
    col_widths=[2, 2, 5, 3]
)

# ============================================================
# 四、商业模式
# ============================================================
doc.add_heading('四、商业模式', level=1)

doc.add_heading('4.1 盈利模式', level=2)

add_table(
    ['收入来源', '模式', '占比(目标)', '说明'],
    [
        ['SaaS订阅', '月付/年付订阅', '60%', '核心收入，四档定价'],
        ['API调用计费', '按量付费', '25%', '企业客户批量调用'],
        ['增值服务', '项目制', '10%', '合规培训/定制规则/认证辅导'],
        ['规则市场分成', '交易抽成', '5%', '用户贡献规则的销售分成'],
    ],
    col_widths=[3, 2.5, 2, 5]
)

doc.add_heading('4.2 定价策略', level=2)

add_table(
    ['版本', '月价', '年价(8折)', '检测额度', '核心权益'],
    [
        ['免费版', '¥0', '-', '50条/月', '1市场+文本检测+基础报告'],
        ['专业版', '¥299', '¥2,870', '2000条/月', '5市场+AI检测+图片检测+PDF导出+数据飞轮'],
        ['企业版', '¥1,999', '¥19,190', '20000条/月', '专业版+平台对接+巡检+API+团队协作'],
        ['定制版', '¥4,999+', '面议', '无限', '企业版+私有部署+定制规则+专属支持'],
    ],
    col_widths=[2, 1.5, 2, 2.5, 5]
)

add_para('定价逻辑：', bold=True)
add_para('• 免费版：降低试用门槛，获取用户和数据飞轮输入；', indent=True)
add_para('• 专业版：覆盖中小卖家核心需求，¥299/月≈¥0.15/条，远低于人工审查成本(¥10-30/条)；', indent=True)
add_para('• 企业版：平台对接和API是关键差异化，锁定大客户；', indent=True)
add_para('• 定制版：高客单价+高粘性，贡献主要利润。', indent=True)

doc.add_heading('4.3 单位经济模型', level=2)

add_table(
    ['指标', '专业版', '企业版', '说明'],
    [
        ['月客单价(ARPU)', '¥299', '¥1,999', '订阅收入'],
        ['边际成本(COGS)', '¥15', '¥80', 'LLM API+服务器+OCR'],
        ['毛利率', '95%', '96%', 'SaaS典型高毛利'],
        ['获客成本(CAC)', '¥500', '¥5,000', '内容/社群 vs 直销'],
        ['LTV(24个月)', '¥7,176', '¥47,976', 'ARPU×24×留存率'],
        ['LTV/CAC', '14.4', '9.6', '>3为健康'],
    ],
    col_widths=[3.5, 2.5, 2.5, 4]
)

# ============================================================
# 五、运营与增长策略
# ============================================================
doc.add_heading('五、运营与增长策略', level=1)

doc.add_heading('5.1 获客策略', level=2)

add_table(
    ['渠道', '目标客户', '策略', '预期CAC'],
    [
        ['内容营销', '中小卖家', '合规案例/法规解读/检测工具SEO', '¥200-300'],
        ['社群运营', '中小卖家', '亚马逊卖家群/跨境社群/知识星球', '¥100-200'],
        ['平台合作', '大卖/服务商', '与ERP/代运营工具集成', '¥2,000-3,000'],
        ['行业会议', '企业客户', '跨境电商峰会/合规论坛Demo', '¥5,000-8,000'],
        ['产品驱动', '所有客户', '免费版→专业版自然转化', '¥50-100'],
    ],
    col_widths=[2.5, 2.5, 4, 2.5]
)

doc.add_heading('5.2 增长飞轮', level=2)

add_para('CrossGuard的增长飞轮由三个子飞轮耦合驱动：', indent=True)

add_para('飞轮1——数据飞轮（精度飞轮）：', bold=True)
add_para('更多用户 → 更多检测 → 更多反馈（误报/漏报标记）→ 更精准规则 → 更低误报率 → 更好口碑 → 更多用户', indent=True)

add_para('飞轮2——平台飞轮（粘性飞轮）：', bold=True)
add_para('平台对接 → Listing自动拉取 → 巡检告警 → 用户依赖 → 迁移成本↑ → 续费率↑ → 更多平台对接', indent=True)

add_para('飞轮3——生态飞轮（网络飞轮）：', bold=True)
add_para('规则市场 → 用户贡献规则 → 规则覆盖度↑ → 新市场/新类别 → 新用户 → 更多规则贡献', indent=True)

doc.add_heading('5.3 合作生态', level=2)

add_table(
    ['合作方', '合作模式', '价值'],
    [
        ['跨境电商ERP(店小秘/马帮等)', 'API集成', '嵌入ERP工作流，降低使用门槛'],
        ['代运营公司', '渠道分销', '代运营客户批量使用，白标方案'],
        ['律所/合规机构', '内容合作', '法规更新+案例共享+客户推荐'],
        ['电商平台(Amazon/Shopee)', '官方合作', '合规工具推荐+API优先支持'],
        ['高校/研究机构', '产学研', '法规研究+算法优化+人才输送'],
    ],
    col_widths=[4, 2.5, 5]
)

# ============================================================
# 六、团队与组织
# ============================================================
doc.add_heading('六、团队与组织', level=1)

add_table(
    ['角色', '职责', '能力要求', '现状'],
    [
        ['项目负责人/产品', '产品规划+商业拓展', '跨境电商+合规领域经验', '已配置'],
        ['后端工程师', '检测引擎+API+平台对接', 'FastAPI+LLM+法规数据处理', '已配置'],
        ['前端工程师', '用户界面+交互体验', 'React+TypeScript+数据可视化', '已配置'],
        ['法务顾问(兼职)', '法规更新+规则审核', '欧盟/美国/东南亚化妆品法规', '需招募'],
        ['增长运营', '获客+社群+内容', '跨境电商社群运营经验', '需招募'],
    ],
    col_widths=[3, 3.5, 4, 2]
)

add_para('团队扩展计划：产品验证阶段（当前）3人核心团队 → 规模化阶段（6个月）5-7人 → 增长阶段（12个月）10-15人。', indent=True)

# ============================================================
# 七、财务预测
# ============================================================
doc.add_heading('七、财务预测', level=1)

add_table(
    ['指标', '第1年', '第2年', '第3年', '第4年', '第5年'],
    [
        ['付费用户', '200', '1,500', '5,000', '12,000', '25,000'],
        ['ARPU(月)', '¥399', '¥499', '¥599', '¥649', '¥699'],
        ['ARR', '¥96万', '¥898万', '¥3,594万', '¥9,345万', '¥2.1亿'],
        ['毛利率', '92%', '94%', '95%', '96%', '96%'],
        ['运营费用', '¥180万', '¥450万', '¥900万', '¥1,800万', '¥3,000万'],
        ['净利润', '-¥84万', '¥395万', '¥2,514万', '¥7,171万', '¥1.7亿'],
        ['累计现金流', '-¥84万', '¥311万', '¥2,825万', '¥9,996万', '¥2.7亿'],
    ],
    col_widths=[2.5, 2, 2, 2, 2, 2]
)

add_para('关键假设：', bold=True)
add_para('• 第1年以产品验证和种子用户为主，200付费用户×¥399月均价；', indent=True)
add_para('• 第2年产品-市场匹配确认，通过内容营销和社群获客，用户增长7.5倍；', indent=True)
add_para('• 第3年平台对接和数据飞轮效应显现，企业客户占比提升，ARPU上升；', indent=True)
add_para('• 毛利率持续提升：规模效应摊薄LLM API成本，数据飞轮减少人工规则维护成本；', indent=True)
add_para('• 第3年实现盈亏平衡，净利润转正。', indent=True)

add_para('融资需求：', bold=True)
add_para('种子轮：寻求200-500万元融资，用于产品打磨（40%）、市场验证（30%）、团队扩展（20%）、合规储备（10%）。预计12-18个月达到产品-市场匹配。', indent=True)

# ============================================================
# 八、风险分析与应对
# ============================================================
doc.add_heading('八、风险分析与应对', level=1)

add_table(
    ['风险类别', '风险描述', '概率', '影响', '应对策略'],
    [
        ['技术风险', 'LLM API成本上升或服务不稳定', '中', '高', '多后端支持(OpenAI/DeepSeek/Ollama)+本地部署降级+缓存机制'],
        ['技术风险', 'AI检测误报率过高导致用户流失', '中', '高', '数据飞轮持续优化+置信度阈值调优+人工审核兜底'],
        ['市场风险', '竞对模仿核心功能', '高', '中', '加速数据飞轮积累+平台对接锁定+法规知识库持续扩充'],
        ['市场风险', '目标市场法规重大变更', '低', '高', '法务顾问持续跟踪+规则热更新+用户通知机制'],
        ['合规风险', '平台API政策收紧', '中', '中', '多平台分散依赖+官方合作争取+手动导入兜底'],
        ['运营风险', '获客成本高于预期', '中', '中', '产品驱动增长(PLG)+免费版降低门槛+社群口碑传播'],
        ['财务风险', '现金流断裂', '低', '极高', '控制烧钱速度+优先验证付费意愿+保持6个月 runway'],
    ],
    col_widths=[2, 3.5, 1, 1, 5]
)

add_para('风险优先级矩阵：', bold=True)
add_para('• 高优先级（概率×影响=高）：AI误报率、LLM成本/稳定性 → 通过数据飞轮和多后端架构已部分缓解；', indent=True)
add_para('• 中优先级：竞对模仿、获客成本 → 通过护城河建设和PLG策略应对；', indent=True)
add_para('• 低优先级但高影响：法规变更、现金流 → 通过法务跟踪和财务纪律防范。', indent=True)

# ============================================================
# 九、社会价值
# ============================================================
doc.add_heading('九、社会价值', level=1)

add_para('CrossGuard不仅具有商业价值，更承载重要的社会使命：', indent=True)

add_para('1. 保护消费者权益', bold=True)
add_para('通过自动识别虚假广告和医疗宣称，防止消费者被误导购买无效甚至有害的产品，减少因违规产品导致的健康损害。', indent=True)

add_para('2. 助力中国品牌合规出海', bold=True)
add_para('降低合规专业门槛，让中小卖家也能遵守目标市场法规，提升"中国制造"的国际形象和信任度，从"低价竞争"走向"合规竞争"。', indent=True)

add_para('3. 促进跨境电商行业健康发展', bold=True)
add_para('通过数据飞轮积累的合规知识，推动行业合规标准提升，减少"劣币驱逐良币"现象，让合规经营者获得竞争优势。', indent=True)

add_para('4. 支持东南亚等新兴市场监管建设', bold=True)
add_para('CrossGuard对东盟ACD等新兴法规体系的支持，帮助东南亚市场监管从"事后处罚"转向"事前预防"，降低社会监管成本。', indent=True)

add_para('5. 知识普惠', bold=True)
add_para('免费版提供50条/月免费检测，让资金有限的初创卖家也能获得专业合规审查服务，促进合规知识普惠。', indent=True)

# ============================================================
# 十、附录
# ============================================================
doc.add_heading('十、附录', level=1)

doc.add_heading('附录A：法规知识库清单', level=2)

add_table(
    ['数据类型', '市场', '文件', '条目数'],
    [
        ['医疗宣称禁用词', 'EU', 'eu_cosmetics_medical.txt', '167词'],
        ['医疗宣称禁用词', 'US', 'us_cosmetics_medical.txt', '122词'],
        ['医疗宣称禁用词', 'SEA', 'sea_cosmetics_medical.txt', '93词'],
        ['绝对化用语', '通用', 'absolute_terms.txt', '75词'],
        ['禁用成分', 'EU', 'eu_cosmetics_ingredients.txt', '35种'],
        ['禁用成分', 'US', 'us_cosmetics_ingredients.txt', '31种'],
        ['禁用成分', 'SEA', 'sea_cosmetics_ingredients.txt', '43种'],
        ['法规数据', 'EU', 'eu_cosmetics.json', '3部法规'],
        ['法规数据', 'US', 'us_cosmetics.json', '4部法规'],
        ['法规数据', 'SEA', 'sea_cosmetics.json', '4部法规'],
        ['替换映射', 'EU', 'eu_cosmetics.json', '30组'],
        ['替换映射', 'US', 'us_cosmetics.json', '20组'],
        ['替换映射', 'SEA', 'sea_cosmetics.json', '26组'],
        ['违规案例', '通用', 'violations.json', '4案例'],
    ],
    col_widths=[3, 2, 4.5, 2]
)

doc.add_heading('附录B：API端点清单', level=2)

add_table(
    ['方法', '端点', '功能', '认证'],
    [
        ['POST', '/api/v1/check', '合规检测', '需认证'],
        ['POST', '/api/v1/check/batch', '批量检测', '需认证'],
        ['POST', '/api/v1/check/image', '图片检测', '需认证'],
        ['GET', '/api/v1/reports', '报告列表', '公开'],
        ['GET', '/api/v1/reports/{id}', '报告详情', '公开'],
        ['DELETE', '/api/v1/reports/{id}', '删除报告', '需admin'],
        ['GET', '/api/v1/reports/{id}/export/pdf', '导出PDF', '需认证'],
        ['GET', '/api/v1/markets', '市场列表', '公开'],
        ['GET', '/api/v1/markets/{m}/categories', '类别列表', '公开'],
        ['GET', '/api/v1/labels', '必需标签', '公开'],
        ['GET', '/api/v1/certifications', '必需认证', '公开'],
        ['POST', '/api/v1/auth/login', '用户登录', '公开'],
        ['POST', '/api/v1/auth/register', '用户注册', '公开'],
        ['GET', '/api/v1/auth/me', '当前用户', '需认证'],
        ['POST', '/api/v1/feedback', '提交反馈', '需认证'],
        ['GET', '/api/v1/feedback/accuracy', '精度指标', '需认证'],
        ['GET', '/api/v1/feedback/suggestions', '优化建议', '需认证'],
        ['GET', '/api/v1/feedback/list', '反馈列表', '需认证'],
        ['GET', '/api/v1/platforms', '平台状态', '需认证'],
        ['POST', '/api/v1/patrol', '触发巡检', '需admin'],
        ['GET', '/api/v1/patrol/history', '巡检历史', '需认证'],
    ],
    col_widths=[1.5, 4.5, 3, 2]
)

doc.add_heading('附录C：技术栈版本', level=2)

add_table(
    ['组件', '版本', '用途'],
    [
        ['Python', '3.12', '后端运行时'],
        ['FastAPI', '0.104+', 'Web框架'],
        ['SQLAlchemy', '2.0+', 'ORM'],
        ['PostgreSQL', '16', '生产数据库'],
        ['React', '19', '前端框架'],
        ['TypeScript', '6.0', '类型安全'],
        ['Ant Design', '6', 'UI组件库'],
        ['Vite', '8', '构建工具'],
        ['Docker', '-', '容器化部署'],
    ],
    col_widths=[3, 2, 5]
)

# ===== 保存 =====
output_path = r'e:\WORKS\LAW\crossguard\docs\出海法盾_CrossGuard_商业计划书.docx'
doc.save(output_path)
print(f'商业计划书已生成: {output_path}')
